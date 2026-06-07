from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "out"
IMAGE_DIR = ROOT / "images"
DOCX_PATH = OUT_DIR / "verl_profiler_partial_token_design.docx"

PAGE_WIDTH_IN = 6.5
COLORS = {
    "ink": (23, 43, 77),
    "blue": (46, 116, 181),
    "light_blue": (233, 242, 250),
    "lighter_blue": (245, 249, 253),
    "border": (172, 190, 210),
    "dark": (38, 45, 55),
    "gray": (95, 105, 117),
    "light_gray": (242, 244, 247),
    "green": (46, 125, 50),
    "orange": (230, 139, 43),
    "red": (183, 47, 47),
    "white": (255, 255, 255),
    "black": (0, 0, 0),
}


@dataclass
class FigureAsset:
    filename: str
    caption: str
    width_in: float

    @property
    def path(self) -> Path:
        return IMAGE_DIR / self.filename


FIGURES = [
    FigureAsset(
        filename="figure1_architecture.png",
        caption="图1 Verl 训推分进程场景下的 Profiler 控制与采集架构",
        width_in=6.2,
    ),
    FigureAsset(
        filename="figure2_param_flow.png",
        caption="图2 部分 Token 采集的配置建模与后端参数映射流程",
        width_in=6.2,
    ),
    FigureAsset(
        filename="figure3_sequence.png",
        caption="图3 一次带部分 Token 采集的 Rollout Step 时序图",
        width_in=6.2,
    ),
]


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)


def font_candidates() -> list[Path]:
    windir = Path("C:/Windows/Fonts")
    return [
        windir / "msyh.ttc",
        windir / "msyhbd.ttc",
        windir / "simhei.ttf",
        windir / "arial.ttf",
    ]


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = font_candidates()
    if bold:
        candidates = [Path("C:/Windows/Fonts/msyhbd.ttc"), *candidates]
    for path in candidates:
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def draw_rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: tuple[int, int, int], outline: tuple[int, int, int], radius: int = 20) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)


def center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.ImageFont, fill: tuple[int, int, int], spacing: int = 6) -> None:
    left, top, right, bottom = box
    max_width = right - left - 24
    lines: list[str] = []
    for raw_line in text.split("\n"):
        current = ""
        for ch in raw_line:
            trial = current + ch
            if draw.textlength(trial, font=font) > max_width and current:
                lines.append(current)
                current = ch
            else:
                current = trial
        if current:
            lines.append(current)
    line_height = font.size + spacing
    total_height = line_height * len(lines) - spacing
    y = top + (bottom - top - total_height) / 2
    for line in lines:
        width = draw.textlength(line, font=font)
        x = left + (right - left - width) / 2
        draw.text((x, y), line, fill=fill, font=font)
        y += line_height


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: tuple[int, int, int], width: int = 5, label: str | None = None, label_offset: tuple[int, int] = (0, 0)) -> None:
    draw.line([start, end], fill=fill, width=width)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    if dx == 0 and dy == 0:
        return
    import math

    angle = math.atan2(dy, dx)
    arrow_len = 18
    wing = 8
    p1 = (
        end[0] - arrow_len * math.cos(angle) + wing * math.sin(angle),
        end[1] - arrow_len * math.sin(angle) - wing * math.cos(angle),
    )
    p2 = (
        end[0] - arrow_len * math.cos(angle) - wing * math.sin(angle),
        end[1] - arrow_len * math.sin(angle) + wing * math.cos(angle),
    )
    draw.polygon([end, p1, p2], fill=fill)
    if label:
        font = load_font(24)
        mid = ((start[0] + end[0]) / 2 + label_offset[0], (start[1] + end[1]) / 2 + label_offset[1])
        label_box = draw.textbbox((0, 0), label, font=font)
        label_w = label_box[2] - label_box[0]
        label_h = label_box[3] - label_box[1]
        padding_x = 12
        padding_y = 6
        rect = (
            int(mid[0] - label_w / 2 - padding_x),
            int(mid[1] - label_h / 2 - padding_y),
            int(mid[0] + label_w / 2 + padding_x),
            int(mid[1] + label_h / 2 + padding_y),
        )
        draw.rounded_rectangle(rect, radius=10, fill=(255, 255, 255), outline=COLORS["border"], width=2)
        draw.text((rect[0] + padding_x, rect[1] + padding_y - 2), label, font=font, fill=COLORS["ink"])


def create_architecture_figure(path: Path) -> None:
    image = Image.new("RGB", (1880, 1080), COLORS["white"])
    draw = ImageDraw.Draw(image)
    title_font = load_font(44, bold=True)
    box_font = load_font(26)
    box_font_bold = load_font(30, bold=True)
    note_font = load_font(24)

    draw.text((70, 44), "Verl 训推分进程下的 Profiler 控制与采集架构", fill=COLORS["ink"], font=title_font)

    left_box = (80, 180, 470, 470)
    manager_box = (560, 180, 1010, 470)
    replica_box = (1100, 120, 1760, 585)
    output_box = (1180, 720, 1720, 970)

    draw_rounded_box(draw, left_box, COLORS["light_blue"], COLORS["blue"])
    center_text(
        draw,
        left_box,
        "Ray Trainer / PPO 主控\n\n按 global step 判断\n是否开启当前轮采集\n\n训练侧继续沿用装饰器采集",
        box_font,
        COLORS["dark"],
    )
    draw.text((120, 142), "训练侧控制面", fill=COLORS["blue"], font=box_font_bold)

    draw_rounded_box(draw, manager_box, (248, 250, 252), COLORS["border"])
    center_text(
        draw,
        manager_box,
        "LLMServerManager\n\nstart_profile()/stop_profile()\n并发 fan-out 到各 rollout replica",
        box_font,
        COLORS["dark"],
    )
    draw.text((642, 142), "Rollout 管理面", fill=COLORS["blue"], font=box_font_bold)

    draw_rounded_rectangle = draw_rounded_box
    draw_rounded_rectangle(draw, replica_box, COLORS["lighter_blue"], COLORS["border"])
    draw.text((1140, 82), "推理侧 Replica 与后端引擎", fill=COLORS["blue"], font=box_font_bold)
    lane1 = (1140, 165, 1720, 285)
    lane2 = (1140, 325, 1720, 445)
    lane3 = (1140, 485, 1720, 605)
    for lane, title, body in [
        (lane1, "Replica.start_profile()", "DistProfiler 先做 enable / rank / discrete 判断"),
        (lane2, "后端参数适配", "vLLM: delay_iterations / max_iterations\nSGLang: start_step / num_steps"),
        (lane3, "Engine 侧执行", "仅采集 response token 窗口 [start, end)\n未配置时回退为全量 rollout 采集"),
    ]:
        draw_rounded_box(draw, lane, COLORS["white"], COLORS["border"], radius=16)
        draw.text((lane[0] + 24, lane[1] + 14), title, fill=COLORS["ink"], font=box_font_bold)
        center_text(draw, (lane[0] + 16, lane[1] + 44, lane[2] - 16, lane[3] - 10), body, load_font(24), COLORS["dark"])

    draw_rounded_box(draw, output_box, (250, 248, 243), COLORS["orange"])
    center_text(
        draw,
        output_box,
        "Profiler 输出目录\nsave_path/\nagent_loop_rollout_replica_{rank}\n\n便于按 replica 定位引擎侧采集结果",
        box_font,
        COLORS["dark"],
    )

    draw_arrow(draw, (470, 325), (560, 325), COLORS["blue"], label="当前 step 需要采集")
    draw_arrow(draw, (1010, 325), (1100, 325), COLORS["blue"], label="广播 start/stop_profile")
    draw_arrow(draw, (1450, 605), (1450, 720), COLORS["orange"], label="写出采集文件", label_offset=(150, 0))

    note_box = (82, 610, 1010, 960)
    draw_rounded_box(draw, note_box, (253, 252, 249), COLORS["border"])
    draw.text((118, 606), "本次设计聚焦点", fill=COLORS["ink"], font=box_font_bold)
    notes = [
        "1. 新增 profile_token_start / profile_token_end，仅在 rollout role 下生效。",
        "2. token 窗口控制的是推理引擎 decode 阶段的 response token 采集范围。",
        "3. 训练侧 profiler 控制逻辑不改调用方式，推理侧通过服务化参数承接该能力。",
        "4. vllm-ascend 配套适配后，可在 Ascend 推理引擎侧复用同一套窗口语义。",
    ]
    y = 706
    for note in notes:
        draw.text((126, y), note, fill=COLORS["dark"], font=note_font)
        y += 46

    image.save(path)


def create_param_flow_figure(path: Path) -> None:
    image = Image.new("RGB", (1800, 980), COLORS["white"])
    draw = ImageDraw.Draw(image)
    title_font = load_font(44, bold=True)
    box_font = load_font(28)
    strong_font = load_font(30, bold=True)
    small_font = load_font(24)

    draw.text((70, 44), "部分 Token 采集的配置建模与后端参数映射", fill=COLORS["ink"], font=title_font)

    top_boxes = [
        ((80, 160, 480, 345), COLORS["light_blue"], "统一配置入口", "tool_config.{torch|npu}\n新增 profile_token_start / profile_token_end"),
        ((540, 160, 940, 345), (245, 248, 252), "配置校验", "类型为 int 或 null\n非负\n若 start/end 同时存在，则 end > start"),
        ((1000, 160, 1400, 345), COLORS["lighter_blue"], "语义约束", "仅 rollout role 使用\nnull/null 表示全量采集\n运行期再与 response 长度匹配"),
    ]
    for box, fill, title, body in top_boxes:
        draw_rounded_box(draw, box, fill, COLORS["border"])
        draw.text((box[0] + 24, box[1] + 20), title, fill=COLORS["ink"], font=strong_font)
        center_text(draw, (box[0] + 16, box[1] + 56, box[2] - 16, box[3] - 16), body, box_font, COLORS["dark"])

    draw_arrow(draw, (480, 252), (540, 252), COLORS["blue"])
    draw_arrow(draw, (940, 252), (1000, 252), COLORS["blue"])

    vllm_box = (140, 470, 800, 780)
    sglang_box = (1000, 470, 1660, 780)
    draw_rounded_box(draw, vllm_box, (247, 251, 255), COLORS["blue"])
    draw_rounded_box(draw, sglang_box, (249, 250, 252), COLORS["border"])

    draw.text((182, 430), "vLLM / vLLM-Ascend 路径", fill=COLORS["blue"], font=strong_font)
    draw.text((1040, 430), "SGLang 路径", fill=COLORS["blue"], font=strong_font)

    vllm_lines = [
        "build_vllm_profiler_args()",
        "delay_iterations = profile_token_start or 0",
        "max_iterations = profile_token_end - delay_iterations or 0",
        "ignore_frontend = true",
        "vllm-ascend 配套适配后复用同一套 profiler_config 语义",
    ]
    sglang_lines = [
        "build_sglang_profiler_args()",
        "start_step = profile_token_start",
        "num_steps = profile_token_end - profile_token_start",
        "未配置窗口时 start_step / num_steps 均为 None",
        "通过 tokenizer_manager.start_profile(**profile_args) 下发",
    ]
    y = 492
    for line in vllm_lines:
        draw.text((182, y), f"- {line}", fill=COLORS["dark"], font=box_font)
        y += 52
    y = 492
    for line in sglang_lines:
        draw.text((1040, y), f"- {line}", fill=COLORS["dark"], font=box_font)
        y += 52

    foot_box = (140, 830, 1660, 920)
    draw_rounded_box(draw, foot_box, (252, 249, 243), COLORS["orange"])
    foot_text = (
        "窗口采集的收益是把 rollout profiler 关注范围缩到目标 decode token 区间，\n"
        "减少全量采集的数据量与噪声，同时保持训练侧控制方式不变。"
    )
    center_text(draw, foot_box, foot_text, small_font, COLORS["dark"])

    image.save(path)


def create_sequence_figure(path: Path) -> None:
    image = Image.new("RGB", (1800, 1400), COLORS["white"])
    draw = ImageDraw.Draw(image)
    title_font = load_font(44, bold=True)
    lane_font = load_font(28, bold=True)
    text_font = load_font(24)
    mini_font = load_font(22)

    draw.text((70, 44), "一次带部分 Token 采集的 Rollout Step 时序", fill=COLORS["ink"], font=title_font)

    lanes = [
        ("RayTrainer", 180),
        ("LLMServerManager", 600),
        ("RolloutReplica", 1020),
        ("Backend Engine", 1440),
    ]
    for label, x in lanes:
        draw.text((x - 90, 128), label, fill=COLORS["blue"], font=lane_font)
        draw.line([(x, 170), (x, 1160)], fill=COLORS["border"], width=4)

    steps = [
        (230, 180, 600, "判断 curr_step_profile=True"),
        (340, 180, 600, "start_profile()"),
        (450, 600, 1020, "并发 fan-out"),
        (560, 1020, 1440, "build profile args\n并调用 engine.start_profile"),
        (670, 180, 1020, "generate_sequences()"),
        (780, 1020, 1440, "response token 0..N 依次 decode"),
        (900, 1440, 1440, "仅采集 [start, end)\n例如 [12, 46)"),
        (1060, 180, 600, "stop_profile()"),
        (1170, 600, 1020, "广播 stop_profile"),
        (1280, 1020, 1440, "写出结果目录"),
    ]

    for y, start_x, end_x, label in steps:
        if start_x == end_x:
            box = (start_x - 120, y - 42, start_x + 120, y + 42)
            draw_rounded_box(draw, box, (248, 250, 252), COLORS["border"], radius=18)
            center_text(draw, box, label, mini_font, COLORS["dark"], spacing=4)
        else:
            draw_arrow(draw, (start_x, y), (end_x, y), COLORS["blue"], width=4)
            label_box = (min(start_x, end_x) + 40, y - 36, max(start_x, end_x) - 40, y + 36)
            center_text(draw, label_box, label, text_font, COLORS["dark"], spacing=4)

    timeline_box = (1040, 760, 1720, 1080)
    draw_rounded_box(draw, timeline_box, (254, 249, 243), COLORS["orange"])
    draw.text((1080, 648), "Token 窗口示意", fill=COLORS["ink"], font=lane_font)
    bar_x0 = 1120
    bar_y = 885
    bar_w = 520
    draw.rounded_rectangle((bar_x0, bar_y, bar_x0 + bar_w, bar_y + 42), radius=16, fill=COLORS["light_gray"], outline=COLORS["border"], width=2)
    start_ratio = 12 / 64
    end_ratio = 46 / 64
    active_x0 = bar_x0 + int(bar_w * start_ratio)
    active_x1 = bar_x0 + int(bar_w * end_ratio)
    draw.rounded_rectangle((active_x0, bar_y, active_x1, bar_y + 42), radius=16, fill=COLORS["orange"], outline=COLORS["orange"], width=2)
    draw.text((bar_x0 - 16, bar_y + 62), "token 0", fill=COLORS["gray"], font=mini_font)
    draw.text((active_x0 - 24, bar_y - 46), "start=12", fill=COLORS["orange"], font=mini_font)
    draw.text((active_x1 - 12, bar_y - 46), "end=46", fill=COLORS["orange"], font=mini_font)
    draw.text((bar_x0 + bar_w - 60, bar_y + 62), "token N", fill=COLORS["gray"], font=mini_font)
    note = "橙色区间为 profiler 真正采集的 response token 范围；\n前后 token 继续正常生成，但不进入目标采集窗口。"
    center_text(draw, (1070, 980, 1690, 1045), note, mini_font, COLORS["dark"])

    image.save(path)


def set_doc_language(run) -> None:
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")


def style_run(run, size: float = 11, bold: bool = False, color: tuple[int, int, int] = COLORS["dark"], name: str = "Microsoft YaHei") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    set_doc_language(run)


def set_cell_background(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_column_widths(table, widths_in: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)


def style_table(table, header_fill: str = "F2F4F7", font_size: float = 10.5) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.12
                for run in p.runs:
                    style_run(run, size=font_size, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_background(cell, header_fill)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    style_run(run, size=10.5, color=COLORS["gray"])


def add_normal_paragraph(doc: Document, text: str, bold: bool = False, color: tuple[int, int, int] = COLORS["dark"]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    style_run(run, size=11, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_paragraph()
    heading.style = f"Heading {level}"
    run = heading.add_run(text)
    if level == 1:
        style_run(run, size=16, bold=True, color=COLORS["blue"])
        heading.paragraph_format.space_before = Pt(16)
        heading.paragraph_format.space_after = Pt(8)
    elif level == 2:
        style_run(run, size=13, bold=True, color=COLORS["blue"])
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    else:
        style_run(run, size=11.5, bold=True, color=COLORS["ink"])
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(4)


def add_note_box(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.rows[0].cells[0].width = Inches(PAGE_WIDTH_IN)
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "F4F6F9")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run, size=10.5, color=COLORS["ink"])
    doc.add_paragraph()


def add_figure(doc: Document, figure: FigureAsset) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    run.add_picture(str(figure.path), width=Inches(figure.width_in))
    add_caption(doc, figure.caption)


def build_document() -> None:
    ensure_dirs()
    create_architecture_figure(FIGURES[0].path)
    create_param_flow_figure(FIGURES[1].path)
    create_sequence_figure(FIGURES[2].path)

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("模块设计说明：Verl Profiler 支持部分 Token 采集")
    style_run(run, size=22, bold=True, color=COLORS["ink"])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("基于 commit dfe0a1edbc6d6a473c2088097c984ef94aa8cd45")
    style_run(run, size=10.5, color=COLORS["gray"])

    add_note_box(
        doc,
        "说明范围：本文仅覆盖 rollout 场景下 profiler 的部分 token 采集能力、后端参数适配与测试设计；"
        "commit 中 torch_memory 相关代码迁移不在本次设计说明范围内。",
    )

    add_heading(doc, "1 Story概述（必要）", 1)
    overview = doc.add_table(rows=2, cols=5)
    overview.autofit = False
    overview.cell(0, 0).text = "包需求名称"
    overview.cell(0, 1).text = "设计需求名称"
    overview.cell(0, 2).text = "story名称"
    overview.cell(0, 3).text = "story描述"
    overview.cell(0, 4).text = "是否做Story设计"
    overview.cell(1, 0).text = "AR20260120839351"
    overview.cell(1, 1).text = "【msprof】支持RL场景训推分进程下按部分token采集"
    overview.cell(1, 2).text = "verl profiler system supports partial token"
    overview.cell(1, 3).text = (
        "在 Verl 训推分进程的 rollout 场景下，为推理侧 profiler 增加 response token 窗口化采集能力，"
        "支持只采集指定 decode token 区间，并将统一配置透传到 vLLM / SGLang 后端；"
        "配套说明 vllm-ascend 侧的服务化参数适配。"
    )
    overview.cell(1, 4).text = "是"
    set_table_column_widths(overview, [1.15, 1.45, 1.25, 2.05, 0.6])
    style_table(overview)

    req_table = doc.add_table(rows=4, cols=2)
    req_table.autofit = False
    req_table.cell(0, 0).text = "需求项"
    req_table.cell(0, 1).text = "说明"
    req_table.cell(1, 0).text = "输入"
    req_table.cell(1, 1).text = "Verl 拉起 RL 训练，采用训推分进程模式，使能 profiler，并在 rollout 角色的 torch / npu 配置中设置可选的 profile_token_start / profile_token_end。"
    req_table.cell(2, 0).text = "处理过程"
    req_table.cell(2, 1).text = "训练主控按 step 决定是否启停 profiler；rollout 管理器将 start_profile / stop_profile 广播到各 replica；后端引擎仅对 response token 窗口 [start, end) 进行采集。"
    req_table.cell(3, 0).text = "输出"
    req_table.cell(3, 1).text = "推理阶段 profiler 在指定的 token 区间内按配置正常采集，并按 replica 输出到 save_path/agent_loop_rollout_replica_{rank}。"
    set_table_column_widths(req_table, [1.1, 5.4])
    style_table(req_table)

    add_normal_paragraph(
        doc,
        "【需求背景】随着大模型强化学习规模化落地，训推分进程架构已成为 Verl rollout 的常见部署形态。"
        "在该模式下，训练侧与推理侧 profiler 的控制链路已经解耦，但原有 rollout 侧采集仍以整个生成阶段为基本粒度，"
        "会带来数据量大、无关 token 噪声多、定位目标 decode 区间困难等问题。"
        "本次设计通过在 rollout role 下新增 response token 窗口配置，把采集范围缩小到用户关注的部分 token，"
        "从而让 RL 场景下的推理 profiling 更可控、更聚焦，也为 Ascend 场景的配套适配提供统一语义。",
    )

    add_figure(doc, FIGURES[0])

    add_heading(doc, "2 Story上下文（必要）", 1)
    add_normal_paragraph(doc, "不涉及新增业务上下文；本次属于既有 rollout profiler 服务化能力上的细化增强。")

    add_heading(doc, "3 功能点分解（必要）", 1)
    fp_table = doc.add_table(rows=5, cols=3)
    fp_table.autofit = False
    fp_table.cell(0, 0).text = "序号"
    fp_table.cell(0, 1).text = "功能点名称"
    fp_table.cell(0, 2).text = "功能点描述"
    entries = [
        ("1", "Rollout 侧部分 Token 窗口化采集", "在 torch / npu profiler tool_config 中新增 profile_token_start、profile_token_end，仅对 rollout role 的 response token 生效，未配置时保持全量 rollout 采集语义不变。"),
        ("2", "统一配置建模与参数校验", "在统一 profiler 配置模型中完成类型、非负值、start/end 前后关系校验，保证窗口配置进入服务端前具备基本合法性。"),
        ("3", "vLLM / SGLang 服务化参数适配", "将统一窗口语义翻译为后端服务化参数：vLLM 使用 delay_iterations / max_iterations，SGLang 使用 start_step / num_steps，并继续沿用 replica 级输出目录语义。"),
        ("4", "vLLM-Ascend 配套适配说明", "作为配套功能点，在 vllm-ascend 侧承接同一组 profiler 窗口参数语义，使 Ascend 推理引擎能够按部分 token 窗口采集（参考用户提供的 vllm-project/vllm-ascend#8953）。"),
    ]
    for row_idx, entry in enumerate(entries, start=1):
        for col_idx, value in enumerate(entry):
            fp_table.cell(row_idx, col_idx).text = value
    set_table_column_widths(fp_table, [0.55, 1.7, 4.25])
    style_table(fp_table)

    add_heading(doc, "4 实现设计（必要）", 1)
    add_heading(doc, "4.1 功能实现思路（必要）", 2)
    add_normal_paragraph(
        doc,
        "1. 窗口化采集能力建模：在 TorchProfilerToolConfig 和 NPUToolConfig 中新增 profile_token_start / profile_token_end，"
        "语义定义为 response token 的左闭右开区间 [start, end)。null 表示不设置边界，从而保留对整个 rollout 阶段的全量采集。"
        "配置侧先做基础校验，运行期再由后端结合实际 response 长度判断是否生效。",
    )
    add_normal_paragraph(
        doc,
        "2. 训练侧控制链路不变：trainer 仍按 global step 决定是否启停 profiler。"
        "当当前 step 命中 profiling 条件时，PPO trainer 调用 llm_server_manager.start_profile()，"
        "生成结束后再调用 stop_profile()；这使部分 token 采集成为 rollout 侧的一种更细粒度采集模式，而不是另一条独立控制流。",
    )
    add_normal_paragraph(
        doc,
        "3. 推理后端承接窗口语义：vLLM 路径在引擎初始化时通过 build_vllm_profiler_args() 注入 profiler_config，"
        "并把 start/end 转换为 delay_iterations / max_iterations；SGLang 路径在离散模式下调用 tokenizer_manager.start_profile(**profile_args)，"
        "对应地把窗口语义转换为 start_step / num_steps。这样，统一配置能够在不同后端下保持相同的用户心智。",
    )
    add_normal_paragraph(
        doc,
        "4. vLLM-Ascend 配套适配：针对 Ascend 推理引擎，配套在 vllm-ascend 侧补齐对服务化 profiler 窗口参数的承接，"
        "使 NPU 场景下也能复用 delay_iterations / max_iterations 这一组语义，避免 Verl 侧为不同设备后端维护两套窗口配置规则。",
    )

    add_figure(doc, FIGURES[1])

    add_heading(doc, "4.1.1 流程图", 3)
    add_figure(doc, FIGURES[2])

    add_heading(doc, "4.1.2 流程说明", 3)
    flow_table = doc.add_table(rows=7, cols=2)
    flow_table.autofit = False
    flow_table.cell(0, 0).text = "阶段"
    flow_table.cell(0, 1).text = "说明"
    flow_rows = [
        ("配置解析", "从 actor_rollout_ref.actor.profiler.tool_config.{torch|npu} 读取 profile_token_start / profile_token_end；空值表示不对 token 边界做限制。"),
        ("训练侧启停", "trainer 在命中当前 profiling step 时调用 llm_server_manager.start_profile()；生成完成后统一 stop_profile()。"),
        ("Replica 过滤", "每个 rollout replica 在 DistProfiler 内先检查 enable / this_rank / discrete，只有命中的 replica 才真正把 profiler 调用下沉到后端。"),
        ("后端映射", "vLLM 使用 delay_iterations / max_iterations / ignore_frontend；SGLang 使用 start_step / num_steps；Ascend 侧配套适配与 vLLM 语义保持一致。"),
        ("窗口执行", "推理引擎在 decode 过程中仅采集 response token 窗口 [start, end) 的 profiler 数据，窗口外 token 继续正常生成但不进入目标采集。"),
        ("结果落盘", "输出目录仍沿用 save_path/agent_loop_rollout_replica_{rank}，便于按 replica 定位、比对与回溯。"),
    ]
    for row_idx, (stage, desc) in enumerate(flow_rows, start=1):
        flow_table.cell(row_idx, 0).text = stage
        flow_table.cell(row_idx, 1).text = desc
    set_table_column_widths(flow_table, [1.15, 5.35])
    style_table(flow_table)

    add_heading(doc, "4.2 数据库及文件持久化设计（可选）", 2)
    add_normal_paragraph(
        doc,
        "不涉及数据库持久化。文件侧继续沿用 profiler 既有产物目录设计："
        "rollout 副本的 profiler 数据按 save_path/agent_loop_rollout_replica_{rank} 输出。"
        "本次设计仅改变采集窗口，不改变目录语义、命名规则与清理责任。",
    )

    add_heading(doc, "4.3 接口描述（必要）", 2)
    iface_table = doc.add_table(rows=3, cols=2)
    iface_table.autofit = False
    iface_table.cell(0, 0).text = "接口 / 配置项"
    iface_table.cell(0, 1).text = "描述"
    iface_table.cell(1, 0).text = "tool_config.{torch|npu}.profile_token_start"
    iface_table.cell(1, 1).text = "仅对 rollout role 生效；指定 response token 的采集起始下标，0-based，左边界包含。"
    iface_table.cell(2, 0).text = "tool_config.{torch|npu}.profile_token_end"
    iface_table.cell(2, 1).text = "仅对 rollout role 生效；指定 response token 的采集结束下标，0-based，右边界不包含；未设置时采集到 response 结束。"
    set_table_column_widths(iface_table, [2.3, 4.2])
    style_table(iface_table)
    add_normal_paragraph(doc, "此次重构不引入新的外部 API；用户面变化体现在 profiler 配置项扩展，默认值保持向后兼容。")

    add_heading(doc, "4.4 GUI界面（可选）", 2)
    add_normal_paragraph(doc, "不涉及。")

    add_heading(doc, "4.5 代码设计（必要）", 2)
    code_table = doc.add_table(rows=8, cols=2)
    code_table.autofit = False
    code_table.cell(0, 0).text = "代码位置"
    code_table.cell(0, 1).text = "设计要点"
    code_rows = [
        ("verl/utils/profiler/config.py", "新增 TorchProfilerToolConfig / NPUToolConfig 的 profile_token_start、profile_token_end；补充合法性校验；在 build_vllm_profiler_args / build_sglang_profiler_args 中完成后端参数映射。"),
        ("verl/trainer/config/profiler/profiler.yaml", "为 torch / npu profiler 暴露新的窗口配置项，默认值为 null，维持全量采集。"),
        ("verl/trainer/config/rollout/rollout.yaml", "把 rollout 场景下的 profiler 窗口配置透传到 actor_rollout_ref 侧统一配置树。"),
        ("verl/workers/rollout/vllm_rollout/vllm_async_server.py", "在引擎初始化阶段注入 profiler_args；vLLM 0.13.0+ 通过 profiler_config 承接窗口语义。"),
        ("verl/workers/rollout/sglang_rollout/async_sglang_server.py", "在离散模式下构造 profile_args，并通过 tokenizer_manager.start_profile(**profile_args) 把窗口配置下发给 SGLang。"),
        ("tests/utils/test_server_profiler.py", "新增对默认值、Torch/NPU 的 vLLM 参数映射、SGLang 参数映射的单元测试，保证窗口语义翻译正确。"),
        ("外部配套：vllm-project/vllm-ascend#8953", "作为配套能力说明，在 vllm-ascend 侧承接服务化 profiler 窗口参数，保证 Ascend 场景的语义一致性。"),
    ]
    for row_idx, row in enumerate(code_rows, start=1):
        code_table.cell(row_idx, 0).text = row[0]
        code_table.cell(row_idx, 1).text = row[1]
    set_table_column_widths(code_table, [2.35, 4.15])
    style_table(code_table)
    add_normal_paragraph(
        doc,
        "代码实现基线：commit dfe0a1edbc6d6a473c2088097c984ef94aa8cd45（verl profiler system supports partial token）。",
        color=COLORS["gray"],
    )

    add_heading(doc, "5 DFx设计（可选）", 1)
    dfx_table = doc.add_table(rows=5, cols=6)
    dfx_table.autofit = False
    headers = ["功能", "系统元素", "故障模式 / 可能原因", "故障影响", "已有规避", "改进建议"]
    for idx, value in enumerate(headers):
        dfx_table.cell(0, idx).text = value
    dfx_rows = [
        ("部分 token 采集", "用户配置", "start / end 类型错误、负值或 end <= start", "配置不生效，运行前即被拒绝", "配置类 __post_init__ 做基础校验", "继续保持错误信息面向用户可读，必要时在 CLI 层补充示例"),
        ("部分 token 采集", "运行期 response 长度", "窗口超出实际 response 长度", "采集范围与预期不完全一致，可能退化为采集到 response 结束", "文档声明“仅在窗口位于 response 长度内时生效”", "可在后端侧补充更显式的告警与统计"),
        ("后端参数适配", "vLLM / SGLang / vllm-ascend", "后端版本不支持服务化窗口参数", "窗口能力不生效，可能回退为全量采集或按后端默认行为处理", "通过文档与测试说明依赖能力，保持默认全量采集兼容路径", "在版本检查或启动日志中增加能力探测信息"),
        ("结果落盘", "save_path/replica 目录", "采集文件输出失败，如路径权限不足", "当前轮 profiler 数据缺失", "沿用既有 save_path 约束与 replica 独立目录设计", "结合运行日志补充目录写失败提示"),
    ]
    for row_idx, row in enumerate(dfx_rows, start=1):
        for col_idx, value in enumerate(row):
            dfx_table.cell(row_idx, col_idx).text = value
    set_table_column_widths(dfx_table, [0.8, 1.0, 1.55, 1.0, 1.05, 1.1])
    style_table(dfx_table, font_size=9.5)

    add_heading(doc, "5.1 性能设计", 2)
    add_normal_paragraph(
        doc,
        "部分 token 采集能力的核心收益是缩小 rollout profiler 的观测窗口。与对整个生成阶段全量采集相比，"
        "窗口化采集能够减少无关 token 带来的数据量与解析压力，并更聚焦目标 decode 区间。"
        "当 profile_token_start / end 不设置时，系统行为保持不变，不引入额外调度开销。",
    )

    add_heading(doc, "5.2 安全设计", 2)
    sec_table = doc.add_table(rows=11, cols=2)
    sec_table.autofit = False
    sec_table.cell(0, 0).text = "Checklist 内容"
    sec_table.cell(0, 1).text = "是否涉及 / 说明"
    sec_rows = [
        ("1 是否新增输入", "Y。新增 profile_token_start / profile_token_end 配置项。"),
        ("1.1 是否通知资料更新", "Y。同步更新 docs/perf/torch_profiling.md 与 Ascend profiling 中英文文档。"),
        ("1.2 是否对输入设计了安全校验", "Y。对类型、非负值、start/end 前后关系做校验；运行期再由后端结合 response 长度生效。"),
        ("2 是否有跨信任域进程间交互", "N。仍沿用既有 Ray RPC / 服务化 profiler 控制链路，本次不新增跨信任域通信方式。"),
        ("3 是否存在文件操作", "Y。Profiler 结果继续写入 save_path/agent_loop_rollout_replica_{rank}。"),
        ("4 是否涉及网络通信", "N。仅复用现有 rollout 后端服务化接口，不引入新的对外网络面。"),
        ("5 是否涉及注入风险", "N。新增内容为数值配置项，不涉及命令、HTML、YAML 执行路径扩展。"),
        ("6 是否引入第三方库", "N。Verl 侧实现不引入新增依赖。"),
        ("7 是否新增二进制交付件", "N。无新增二进制。"),
        ("8-10 加密/敏感信息/安全函数库", "N / NA。本次仅涉及 profiler 采集窗口控制。"),
    ]
    for row_idx, row in enumerate(sec_rows, start=1):
        sec_table.cell(row_idx, 0).text = row[0]
        sec_table.cell(row_idx, 1).text = row[1]
    set_table_column_widths(sec_table, [2.5, 4.0])
    style_table(sec_table)

    add_heading(doc, "5.3 兼容性设计", 2)
    add_normal_paragraph(
        doc,
        "1. 配置兼容性：新增窗口配置默认值为 null，不影响既有 profiler 行为；非 rollout role 读取到这些字段时不参与生效。",
    )
    add_normal_paragraph(
        doc,
        "2. 后端兼容性：vLLM 路径的窗口能力依赖后端支持 profiler_config 中的 delay_iterations / max_iterations；"
        "SGLang 路径依赖 start_profile(start_step, num_steps) 这一组服务化参数；"
        "Ascend 场景通过 vllm-ascend 配套适配保持与社区 vLLM 相同的窗口语义。",
    )
    add_normal_paragraph(
        doc,
        "3. 结果兼容性：输出目录、rank 选择、离散模式与训练侧 step 控制逻辑均保持原有语义，便于现网平滑使用。",
    )

    add_heading(doc, "5.4 全球化", 2)
    add_normal_paragraph(doc, "不涉及。")

    add_heading(doc, "5.5 日志上报、BI，曝光性及可维护性", 2)
    add_normal_paragraph(doc, "日志上报：不涉及新增日志上报面。")
    add_normal_paragraph(doc, "BI 上报：不涉及。")
    add_normal_paragraph(doc, "可维护性：窗口配置集中在统一 profiler 配置模型中，便于后续扩展到更多 rollout 后端。")

    add_heading(doc, "6 测试设计（必要）", 1)
    add_heading(doc, "6.1 单元测试（UT）", 2)
    ut_table = doc.add_table(rows=5, cols=2)
    ut_table.autofit = False
    ut_table.cell(0, 0).text = "测试项"
    ut_table.cell(0, 1).text = "验证内容"
    ut_rows = [
        ("默认 vLLM 参数", "当不设置 profile_token_start / end 时，delay_iterations=0、max_iterations=0，保持全量采集语义。"),
        ("Torch vLLM 窗口映射", "示例 start=12、end=46 时，校验 delay_iterations=12、max_iterations=34。"),
        ("NPU vLLM 窗口映射", "示例 start=5、end=13 时，校验 delay_iterations=5、max_iterations=8。"),
        ("SGLang 窗口映射", "示例 start=7、end=16 时，校验 start_step=7、num_steps=9；未配置时保持 None。"),
    ]
    for row_idx, row in enumerate(ut_rows, start=1):
        ut_table.cell(row_idx, 0).text = row[0]
        ut_table.cell(row_idx, 1).text = row[1]
    set_table_column_widths(ut_table, [2.0, 4.5])
    style_table(ut_table)

    add_heading(doc, "6.2 接口测试", 2)
    add_normal_paragraph(doc, "不涉及对外接口变更；重点验证配置扩展后仍能通过既有 trainer / rollout 控制链路驱动 profiler。")

    add_heading(doc, "6.3 业务场景测试", 2)
    add_normal_paragraph(
        doc,
        "建议在 RL 训推分进程模式下执行以下场景："
        "1) 不设置窗口，验证行为与原有全量采集一致；"
        "2) 设置 [12,46) 窗口，验证仅在指定 decode token 区间生成 rollout profiler 数据；"
        "3) 在 Ascend + vllm-ascend 场景下使用相同配置，验证配套适配后的窗口语义一致。",
    )

    add_heading(doc, "6.4 异常场景测试", 2)
    add_normal_paragraph(
        doc,
        "建议覆盖："
        "1) start/end 非法配置；"
        "2) response 长度短于目标窗口；"
        "3) 后端版本不支持服务化窗口参数；"
        "4) 训练中断或提前 stop_profile 时的数据落盘与目录完整性。",
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
